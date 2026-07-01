from docx import Document


def generate_resume(text, output_path):
    doc = Document()

    doc.add_heading("Tailored Resume", level=1)

    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.save(output_path)